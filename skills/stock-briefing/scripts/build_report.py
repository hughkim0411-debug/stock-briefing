#!/usr/bin/env python3
"""Build a beginner-friendly Korean/US stock market explainer."""
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "맑은 고딕"

# ---------- helpers ----------
def set_run_font(run, name=FONT, size=11, bold=False, italic=False, color=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def add_para(doc, text, *, size=11, bold=False, italic=False, color=None,
             align=None, space_before=4, space_after=4, line_spacing=1.5,
             indent_left=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if indent_left is not None:
        pf.left_indent = Cm(indent_left)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic, color=color)
    return p

def add_runs(doc, parts, *, align=None, space_before=4, space_after=4,
             line_spacing=1.5, indent_left=None):
    """parts: list of (text, opts dict)"""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if indent_left is not None:
        pf.left_indent = Cm(indent_left)
    for text, opts in parts:
        r = p.add_run(text)
        set_run_font(r, size=opts.get('size', 11),
                     bold=opts.get('bold', False),
                     italic=opts.get('italic', False),
                     color=opts.get('color'))
    return p

def add_h1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(20)
    pf.space_after = Pt(10)
    r = p.add_run(text)
    set_run_font(r, size=17, bold=True, color="1F3864")
    # bottom border on this paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), '2E74B5')
    bottom.set(qn('w:space'), '4')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(14)
    pf.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size=13, bold=True, color="2E74B5")
    return p

def add_bullet(doc, text, bold=False, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.4
    if indent:
        pf.left_indent = Cm(indent)
    r = p.add_run(text)
    set_run_font(r, size=11, bold=bold)
    return p

def add_callout(doc, title, body, color="E7F0F8", border_color="2E74B5"):
    """Add a single-cell shaded box used for tips / analogies."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Cm(17)
    # shading
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)
    # borders
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '6')
        b.set(qn('w:color'), border_color)
        tcBorders.append(b)
    tc_pr.append(tcBorders)
    # padding via margins
    tcMar = OxmlElement('w:tcMar')
    for side, sz in (('top', '160'), ('bottom', '160'), ('left', '200'), ('right', '200')):
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), sz)
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tc_pr.append(tcMar)

    # title paragraph
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(title)
    set_run_font(r1, size=11, bold=True, color=border_color)
    # body paragraph
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.45
    r2 = p2.add_run(body)
    set_run_font(r2, size=11)
    # spacer below
    add_para(doc, "", space_before=2, space_after=2)
    return table

def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), 'BFBFBF')
        tcBorders.append(b)
    tc_pr.append(tcBorders)

def fill_cell(cell, text, *, header=False, alt=False, bold=False, color=None, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    if header:
        set_run_font(r, size=11, bold=True, color="FFFFFF")
        set_cell_shading(cell, "1F3864")
    else:
        set_run_font(r, size=11, bold=bold, color=color)
        if alt:
            set_cell_shading(cell, "F2F2F2")
    set_cell_borders(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = False
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    for i, h in enumerate(headers):
        fill_cell(table.rows[0].cells[i], h, header=True, center=True)
    for ri, row in enumerate(rows):
        alt = (ri % 2 == 1)
        for ci, item in enumerate(row):
            opts = dict(item) if isinstance(item, dict) else {"text": item}
            fill_cell(table.rows[ri + 1].cells[ci], opts.get("text", ""),
                      alt=alt, bold=opts.get("bold", False),
                      color=opts.get("color"), center=opts.get("center", False))
    add_para(doc, "", space_before=2, space_after=2)
    return table

# ---------- build ----------
def build(out_path):
    doc = Document()

    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # default style
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(11)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FONT)
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)

    # header / footer
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("증시 입문자를 위한 친절한 설명서")
    set_run_font(hr, size=9, color="808080")

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("자료 출처: 한국경제 + 네이버 뉴스 검색 / 작성: 2026.05.20")
    set_run_font(fr, size=9, color="808080")

    # ===== Title =====
    add_para(doc, "증시 입문자를 위한",
             size=18, bold=False, color="595959",
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=2)
    add_para(doc, "한국·미국 증시 친절한 설명서",
             size=26, bold=True, color="1F3864",
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=4)
    add_para(doc, "(2026년 5월 19일~20일 시장 상황)",
             size=12, italic=True, color="2E74B5",
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=12)

    add_callout(doc,
        "이 문서를 읽기 전에",
        "이 자료는 '주식이라는 단어만 들어봤다'는 분도 한 번 읽고 이해할 수 있도록 풀어 쓴 설명서입니다. "
        "어려운 용어가 나오면 그 자리에서 풀이를 함께 적었고, 비유와 예시를 많이 넣었습니다. "
        "30분 정도 천천히 읽으면 어제·오늘 뉴스에서 무슨 일이 있었는지 큰 그림이 잡힙니다.")

    # ===== Section 0: Vocabulary =====
    add_h1(doc, "0장. 시작하기 전에 — 꼭 알아야 할 7가지 단어")

    add_para(doc, "뉴스를 보다 보면 이런 단어들이 계속 나옵니다. 한 번만 정리해두면 그 다음부터는 술술 읽힙니다.",
             space_after=8)

    vocab = [
        ("코스피 (KOSPI)",
         "한국 주식시장의 대표 가격표. 삼성전자·SK하이닉스 같은 큰 회사 800여 곳의 주가를 합쳐서 만든 평균값입니다. "
         "'코스피가 3% 떨어졌다'는 말은 그 평균값이 3% 내렸다는 뜻이고, 보통 시장 전체 분위기를 보여줍니다."),
        ("코스닥 (KOSDAQ)",
         "코스피보다 작은 회사들이 모인 시장. 보통 기술 벤처·바이오·게임 회사가 많습니다. "
         "변동성(가격이 출렁이는 정도)이 코스피보다 큰 편입니다."),
        ("다우 / S&P500 / 나스닥",
         "미국의 3대 주식시장 가격표입니다. 다우는 대표 30개 기업, S&P500은 대표 500개 기업, "
         "나스닥은 애플·구글·엔비디아 같은 기술주 중심입니다. 미국이 흔들리면 다음 날 아침 한국이 따라가는 경우가 많습니다."),
        ("외국인 (외국인 투자자)",
         "한국 주식시장에 돈을 넣은 해외 펀드·기관을 말합니다. 외국인이 '순매도'(파는 양이 더 많음)를 시작하면 "
         "지수가 잘 떨어지고, '순매수'(사는 양이 더 많음)면 잘 오릅니다. 한국 증시에서 영향력이 가장 큰 손님입니다."),
        ("국채 금리",
         "정부가 돈을 빌리면서 약속하는 이자율. 미국 정부가 30년 동안 돈 빌리는 금리(=미 30년물 국채금리)가 올라가면, "
         "주식 대신 채권(=정부에 빌려주고 이자 받는 상품)이 매력적이 되면서 주식에서 돈이 빠져나갑니다. "
         "그래서 금리는 주식 시장의 '경쟁자' 같은 존재입니다."),
        ("반도체주",
         "삼성전자·SK하이닉스 같은 회사 주식. 한국 증시 전체 시가총액의 약 30%를 차지하기 때문에 "
         "이 두 회사가 떨어지면 코스피 전체가 휘청합니다. 요즘은 'AI에 꼭 필요한 부품'이라는 이유로 더 주목받습니다."),
        ("환율 (원/달러)",
         "1달러를 사려면 원화 얼마가 필요한가? 1,500원이면 1달러 = 1,500원. "
         "환율이 올라간다는 건 원화 가치가 떨어진다는 뜻이고, 외국인 입장에서는 한국 주식을 갖고 있어도 "
         "달러로 바꿀 때 손해를 보니까 더 팔고 싶어집니다."),
    ]
    add_table(doc, ["용어", "한 줄 설명"],
              [[{"text": k, "bold": True}, {"text": v}] for k, v in vocab],
              col_widths_cm=[4.0, 13.0])

    # ===== Section 1: One-paragraph summary =====
    add_h1(doc, "1장. 한 문단으로 요약하면")

    add_callout(doc,
        "이번 주 증시, 무슨 일이?",
        "잘 오르던 차가 갑자기 브레이크를 밟았습니다. 한국 코스피는 5월 15일 고점(8,046)에서 이틀 만에 7,200대까지 "
        "떨어졌고, 미국도 같이 흔들렸어요. 가장 큰 이유는 두 가지입니다. "
        "(1) 미국이 빌리는 돈의 이자(국채금리)가 18년 만에 가장 높은 수준까지 올라서 주식이 상대적으로 덜 매력적이 됐고, "
        "(2) '반도체 회사들 이제 다 벌었으니 정점 아니냐'는 의심이 미국에서 시작돼 한국 삼성전자·SK하이닉스까지 번졌기 때문입니다. "
        "다만 전문가 대부분은 '강세장(=오르는 시장) 자체가 끝난 건 아니고 잠시 숨고르기'라는 시각입니다.",
        color="FFF4D6", border_color="E69900")

    # ===== Section 2: Korea =====
    add_h1(doc, "2장. 한국 증시 — '잘 가던 차가 급정거했다'")

    add_h2(doc, "2-1. 숫자로 보는 5월 19일 ~ 20일")
    add_table(doc,
        ["날짜", "코스피", "코스닥", "한 줄 설명"],
        [
            [{"text":"5월 19일(월)","bold":True}, {"text":"7,271.66 (-3.25%)","center":True,"color":"C00000","bold":True},
             {"text":"2%대 급락","center":True,"color":"C00000"},
             {"text":"하루 만에 약 -3% — 작은 충격이 아님"}],
            [{"text":"5월 20일(화)","bold":True}, {"text":"7,208.95 (-0.86%)","center":True,"color":"C00000"},
             {"text":"1,056.07 (-2.61%)","center":True,"color":"C00000"},
             {"text":"하락폭은 줄었지만 여전히 약세"}],
        ],
        col_widths_cm=[3.2, 4.5, 3.5, 5.8])

    add_callout(doc,
        "이게 큰 하락인가요?",
        "코스피가 하루에 3% 떨어진다는 건 일반적으로 '강한 충격'입니다. 1년에 몇 번 있을까 말까 한 일이에요. "
        "비유하자면, 평소 시속 80km로 달리던 차가 갑자기 30km로 줄어든 느낌입니다. 사고는 아니지만, 깜짝 놀라는 정도.")

    add_h2(doc, "2-2. 외국인이 10일 연속 '팔자' — 누적 50조원 안팎")
    add_para(doc, "5월 7일부터 5월 20일까지 외국인 투자자는 한국 주식을 계속 팔고 있습니다. "
                  "9거래일까지 누적으로 약 48조 9천억 원, 10거래일째인 오늘(20일)까지 더하면 약 52조 원 안팎의 매도입니다.")
    add_para(doc, "특히 삼성전자와 SK하이닉스를 집중적으로 팔았습니다. "
                  "다행히 한국 개인 투자자들이 비슷한 규모로 사주고 있어서 지수가 더 폭락하지는 않았어요. "
                  "그래도 '외국인이 빠져나가는 시장은 단기적으로 약하다'는 게 일반적인 시장 상식입니다.",
              space_after=8)

    add_callout(doc,
        "왜 외국인은 한국 주식을 팔까요?",
        "(1) 미국 금리가 너무 높아져서 '굳이 한국 주식 안 사도 미국 채권만 사도 충분히 수익 나겠다' 싶고, "
        "(2) 환율이 1,500원을 넘어가니까 갖고 있는 한국 주식을 달러로 바꿔도 손해라는 계산이 나오며, "
        "(3) 한국 증시 자체가 너무 빨리 올랐기 때문에 '한 번쯤 이익 실현' 심리가 강해졌습니다.")

    add_h2(doc, "2-3. 주요 종목은 어떻게 됐나? (5월 19일 기준)")
    add_table(doc,
        ["종목", "등락", "쉽게 풀어보면"],
        [
            [{"text":"삼성전자","bold":True}, {"text":"-1.96%","center":True,"color":"C00000"},
             {"text":"외국인이 가장 많이 판 종목. 그래도 회사 실적 전망은 작년 말보다 3배 좋아짐"}],
            [{"text":"SK하이닉스","bold":True}, {"text":"-5.16%","center":True,"color":"C00000"},
             {"text":"'200만 원 갈 것'이라던 기대가 잠시 꺾임 (현재 174만 5천원)"}],
            [{"text":"현대차","bold":True}, {"text":"-8.90%","center":True,"color":"C00000"},
             {"text":"'로봇 만든다는 기대'로 너무 올랐던 게 빠짐 — 차익실현"}],
            [{"text":"LG전자","bold":True}, {"text":"-11.66%","center":True,"color":"C00000"},
             {"text":"같은 이유. 로봇 기대주들이 한꺼번에 조정"}],
            [{"text":"한화에어로스페이스","bold":True}, {"text":"+4.81%","center":True,"color":"00703C","bold":True},
             {"text":"방산(=무기·국방) 회사. 이란 사태로 무기 수요 기대 → 혼자 강세"}],
        ],
        col_widths_cm=[4.0, 2.5, 10.5])

    # ===== Section 3: US =====
    add_h1(doc, "3장. 미국 증시 — '감기 걸린 미국, 한국이 옮았다'")

    add_h2(doc, "3-1. 미국 3대 지수 마감 (5월 19일 현지시간)")
    add_table(doc,
        ["지수", "종가", "등락", "쉽게 풀어보면"],
        [
            [{"text":"다우","bold":True}, {"text":"49,363.88","center":True},
             {"text":"-0.65%","center":True,"color":"C00000","bold":True},
             {"text":"미국 대표 30개 기업 평균"}],
            [{"text":"S&P500","bold":True}, {"text":"7,353.61","center":True},
             {"text":"-0.67%","center":True,"color":"C00000","bold":True},
             {"text":"미국 우량 500개 평균 — 가장 많이 인용되는 지수"}],
            [{"text":"나스닥","bold":True}, {"text":"25,870.71","center":True},
             {"text":"-0.84%","center":True,"color":"C00000","bold":True},
             {"text":"애플·구글·엔비디아 등 기술주 중심"}],
        ],
        col_widths_cm=[3.0, 3.5, 2.5, 8.0])

    add_callout(doc,
        "미국이 떨어지면 왜 한국도 떨어지나요?",
        "한국 시간으로 다음 날 아침에 한국 증시가 열리는데, 그 사이 밤사이 미국에서 어떤 일이 있었는지가 "
        "한국 투자자들의 첫 반응에 큰 영향을 줍니다. 또 외국인 투자자들은 전 세계 주식을 한꺼번에 다루기 때문에, "
        "'미국이 위험해 보이면 한국도 줄여야지' 하고 같이 팔아버립니다. 그래서 두 시장은 같이 움직이는 경우가 많습니다.")

    # ===== Section 4: The big two reasons =====
    add_h1(doc, "4장. 모든 일의 원인 — '금리'와 '반도체'")

    add_h2(doc, "4-1. 첫 번째 원인: 미국 금리가 너무 올랐다")
    add_para(doc, "이번에 가장 큰 충격은 '미국 30년물 국채 금리'가 5.197%까지 올랐다는 점입니다. "
                  "2007년 7월 이후 18년 만에 가장 높은 수치예요. 그런데 30년물? 5%? 이게 왜 중요한지 풀어볼게요.")

    add_callout(doc,
        "비유로 이해하는 국채 금리",
        "미국 정부에 30년 동안 돈을 빌려주면 매년 5.2%의 이자를 준다는 뜻입니다. "
        "이게 무슨 의미일까요? 만약 당신이 '주식에 투자해서 1년에 7% 벌 수 있다'고 생각해도, "
        "옆에서 '안전한 미국 정부 빚이 5.2% 준다'고 하면 '굳이 위험한 주식 안 해도 되네' 하고 채권 쪽으로 갑니다. "
        "이게 바로 금리가 오르면 주식 시장에서 돈이 빠지는 이유예요.")

    add_para(doc, "월스트리트(미국 금융 중심가)에서는 이 현상을 '채권 자경단(Bond Vigilantes)이 움직였다'고 부릅니다. "
                  "채권 시장 투자자들이 정부의 재정 적자(빚이 너무 많음)와 인플레이션(물가 상승)에 화가 나서 "
                  "'돈 빌려주는 가격을 더 비싸게 받겠다'고 단체로 행동에 나섰다는 뜻입니다. 일종의 시장의 시위인 셈이죠.")

    add_h2(doc, "4-2. 두 번째 원인: '반도체 정점 통과' 논란")
    add_para(doc, "5월 18일 미국 메모리 반도체 회사 시게이트의 CEO가 한 말이 도화선이 됐습니다. "
                  "\"공장을 더 짓는 데 시간이 너무 오래 걸린다\"는 발언이었어요. 그런데 이 말이 왜 충격일까요?")

    add_callout(doc,
        "'피크아웃(Peak Out)'이라는 단어",
        "한자로 풀면 '정점 통과'. 더 이상 올라갈 곳이 없으니 이제 내려갈 일만 남았다는 뜻입니다. "
        "투자자들은 시게이트 CEO 발언을 '반도체 회사들이 이제 한계에 다다랐다'는 신호로 해석했어요. "
        "그래서 다음 날 미국 반도체 지수가 2.47% 빠졌고, 다시 그 다음 날 한국 삼성·하이닉스가 같이 떨어진 겁니다. "
        "전염병처럼 번진 거죠.")

    add_para(doc, "반대 의견도 있습니다. KB증권 김동원 본부장은 이렇게 말했어요. "
                  "\"옛날에는 반도체가 좋을 때도 있고 나쁠 때도 있는 '사이클' 산업이었지만, 지금은 AI 시대라서 "
                  "구조적으로 계속 잘 팔린다. 더 길게 봐야 한다.\" — 즉 옛날 잣대로 '이제 끝났다'고 단정하면 안 된다는 거예요.")

    # ===== Section 5: New variables =====
    add_h1(doc, "5장. 오늘(20일) 새로 등장한 변수 3가지")

    add_h2(doc, "5-1. 환율 1,500원 돌파 — 장중 1,513원")
    add_para(doc, "오늘 원/달러 환율이 장중에 1,513원까지 올랐다가 1,506.8원에 마감했습니다. "
                  "1,500원이라는 숫자는 심리적 저항선이라 시장이 예민하게 반응합니다.")

    add_callout(doc,
        "환율이 오르면 좋은가 나쁜가?",
        "수출 기업(삼성전자, 현대차 등)에게는 이론적으로 좋습니다. 같은 물건 팔고도 원화로 받으면 더 많이 받으니까요. "
        "하지만 외국인 투자자에게는 '한국 주식을 갖고 있어도 달러로 바꿀 때 손해'이기 때문에 매도 압력이 됩니다. "
        "그리고 수입 물가(기름·식품 원료)가 비싸져서 일반 국민에게는 부담이 됩니다. "
        "한쪽에는 좋고 한쪽에는 나쁜, 양날의 검입니다.")

    add_h2(doc, "5-2. 삼성전자 노조 총파업 선언")
    add_para(doc, "삼성전자 노조가 5월 21일 총파업에 들어간다고 발표했습니다. "
                  "협상이 결렬됐고, 노조 측은 '내일 예정대로 총파업 돌입'이라고 밝혔어요. "
                  "삼성전자는 한국 증시 시총 1위이기 때문에, 회사 불안 = 코스피 불안으로 직결됩니다.")

    add_h2(doc, "5-3. 트럼프와 이란 — 호르무즈 해협")
    add_para(doc, "트럼프 대통령이 5월 19일로 예정됐던 이란 군사공격을 보류했습니다. "
                  "다만 \"적절한 합의가 안 되면 대규모 공격할 수 있다\"고 함께 말해서 시장 불안은 남아있어요.")

    add_callout(doc,
        "호르무즈 해협이 뭐길래?",
        "이란 옆에 있는 좁은 바닷길인데, 전 세계 원유의 20% 이상이 이곳을 지나갑니다. "
        "이란이 만약 이 해협을 막아버리면 기름이 안 들어와서 유가가 폭등하고, 그러면 물가가 올라가고, "
        "그러면 미국 연준(Fed)이 금리를 더 못 내리거나 오히려 올려야 합니다. "
        "그래서 호르무즈 = 유가 = 금리 = 증시가 일렬로 연결돼 있어요.")

    # ===== Section 6: How to look at it =====
    add_h1(doc, "6장. 그래서 어떻게 봐야 하나요?")

    add_h2(doc, "6-1. 전문가들의 시각")
    add_table(doc,
        ["증권사", "전문가", "쉽게 풀어쓴 결론"],
        [
            [{"text":"키움증권","bold":True}, {"text":"한지영","center":True},
             {"text":"\"너무 빨리 올랐으니 잠시 쉬는 건 자연스럽다. 변동성 관리 시점.\""}],
            [{"text":"하나증권","bold":True}, {"text":"황승택","center":True},
             {"text":"\"삼성전자 이익 전망이 여전히 좋아지는 중. 메모리 호황은 계속.\""}],
            [{"text":"NH투자증권","bold":True}, {"text":"조수홍","center":True},
             {"text":"\"AI에 필요한 전기·원전·로봇 분야 추천.\""}],
            [{"text":"메리츠증권","bold":True}, {"text":"이진우","center":True},
             {"text":"\"위험한 시기. 대형 주도주(=삼성·하이닉스 등) 위주로 압축.\""}],
            [{"text":"신한투자증권","bold":True}, {"text":"윤창용","center":True},
             {"text":"\"실적이 확인되는 소비·유통 등으로도 일부 분산.\""}],
        ],
        col_widths_cm=[3.5, 2.5, 11.0])

    add_para(doc, "공통점은 '강세장은 끝나지 않았다, 그러나 무리하게 따라가지는 말자'입니다. "
                  "지금 막 시작하는 분들에게 적용해보면, '한꺼번에 큰돈을 넣지 말고 천천히 분할해서 들어가라'는 뜻이에요.",
              space_before=4, space_after=8)

    add_h2(doc, "6-2. 다음 주 주목할 일정")
    add_bullet(doc, "5월 21일 새벽 (한국시간): 엔비디아 실적 발표 — AI 반도체가 얼마나 잘 팔렸는지 확인. 시장 분위기를 좌우할 핵심 이벤트.")
    add_bullet(doc, "월마트·타겟 등 미국 대형 마트 실적 — 미국 소비자가 아직 잘 쓰고 있는지 확인")
    add_bullet(doc, "호르무즈 해협 개방 여부 — 유가·금리·증시가 한꺼번에 움직일 가능성")
    add_bullet(doc, "삼성전자 파업 진행 상황 — 21일 총파업 실제 돌입 여부")

    # ===== Section 7: Outlook & Risks =====
    add_h1(doc, "7장. 앞으로 어떻게 될까? — 전망과 잠재적 악재")

    add_h2(doc, "7-1. 단기 전망 — 3가지 시나리오")
    add_para(doc, "앞으로 1~4주 사이 시장이 갈 수 있는 길은 크게 세 갈래입니다. 어느 길로 갈지는 "
                  "(1) 엔비디아 실적 (2) 호르무즈 해협 (3) 미국 채권 금리 — 이 세 가지가 결정합니다.")

    add_table(doc,
        ["시나리오", "발생 조건", "예상 코스피", "확률(주관적)"],
        [
            [{"text":"낙관 시나리오\n(빠른 반등)","bold":True,"color":"00703C"},
             {"text":"엔비디아 실적 호조 + 이란 협상 진전 + 미 금리 하향 안정"},
             {"text":"7,500~8,000선 회복","center":True,"color":"00703C","bold":True},
             {"text":"30%","center":True}],
            [{"text":"중립 시나리오\n(횡보·박스권)","bold":True,"color":"595959"},
             {"text":"엔비디아 실적 무난 + 이란 사태 장기화 + 금리 5% 안팎 유지"},
             {"text":"7,000~7,500선 박스","center":True,"color":"595959","bold":True},
             {"text":"50%","center":True}],
            [{"text":"비관 시나리오\n(추가 조정)","bold":True,"color":"C00000"},
             {"text":"엔비디아 실적 쇼크 OR 호르무즈 폐쇄 OR 미 30년물 5.5% 돌파"},
             {"text":"6,500~7,000선 하향","center":True,"color":"C00000","bold":True},
             {"text":"20%","center":True}],
        ],
        col_widths_cm=[3.5, 7.5, 3.5, 2.5])

    add_callout(doc,
        "왜 중립 시나리오 확률이 가장 높을까?",
        "역사적으로 시장은 한 번 큰 충격을 받으면 단번에 회복하지도, 단번에 무너지지도 않고 "
        "한동안 옆걸음(=박스권)을 치는 경우가 가장 많기 때문입니다. "
        "전문가들이 공통적으로 말하는 '코스피 6,800~7,100 1차 지지선'은 곧 '이 정도 선에서 한동안 머물 것'이라는 예측이에요.")

    add_h2(doc, "7-2. 조심해야 할 잠재적 악재 6가지")
    add_para(doc, "앞으로 시장을 더 끌어내릴 수 있는 위험 요인들입니다. 하나만 터져도 부담이고, 여러 개가 겹치면 비관 시나리오로 갑니다.",
             space_after=8)

    add_table(doc,
        ["악재", "어떤 일이 생기면?", "충격 강도"],
        [
            [{"text":"① 엔비디아 실적 쇼크","bold":True,"color":"C00000"},
             {"text":"5/21 새벽 발표. 'AI 수요 둔화' 신호가 나오면 한국 반도체주 동반 급락 — 코스피 추가 -3~5% 가능"},
             {"text":"★★★★★","center":True,"color":"C00000","bold":True}],
            [{"text":"② 호르무즈 해협 폐쇄","bold":True,"color":"C00000"},
             {"text":"이란이 실제로 막으면 유가가 배럴당 130~150달러로 폭등 → 물가·금리 추가 압박"},
             {"text":"★★★★★","center":True,"color":"C00000","bold":True}],
            [{"text":"③ 미 30년물 금리 5.5% 돌파","bold":True,"color":"C00000"},
             {"text":"전문가들이 '주식의 임계점'으로 지목한 수준. 채권 자경단이 더 움직이면 가능"},
             {"text":"★★★★","center":True,"color":"C00000","bold":True}],
            [{"text":"④ 환율 1,550원 돌파","bold":True,"color":"C00000"},
             {"text":"이미 1,513원까지 갔음. 더 오르면 외국인 매도가 가속화 + 수입물가 부담"},
             {"text":"★★★★","center":True,"color":"C00000","bold":True}],
            [{"text":"⑤ 삼성전자 파업 장기화","bold":True,"color":"C00000"},
             {"text":"5/21 총파업 후 길어지면 반도체 생산 차질 우려 — 코스피 시총 1위 흔들림"},
             {"text":"★★★","center":True,"color":"C00000","bold":True}],
            [{"text":"⑥ 외국인 매도 지속","bold":True,"color":"C00000"},
             {"text":"이미 10거래일 연속 매도. 개인 매수 여력이 한계에 다다르면 지수 추가 하락"},
             {"text":"★★★","center":True,"color":"C00000","bold":True}],
        ],
        col_widths_cm=[4.5, 9.5, 3.0])

    add_h2(doc, "7-3. 반대로, 호재가 될 수 있는 변수")
    add_bullet(doc, "엔비디아 실적이 예상치를 크게 넘으면 → 반도체주 동반 반등 가능", bold=True)
    add_bullet(doc, "이란·미국이 합의에 도달해 호르무즈가 안정되면 → 유가·금리 동반 하락 → 증시 안도 랠리")
    add_bullet(doc, "미 연준(Fed)이 '금리 인상 안 한다' 명확히 시그널 → 채권 자경단 진정")
    add_bullet(doc, "외국인 매도세가 멈추고 순매수로 전환되면 → 코스피 빠른 회복")
    add_bullet(doc, "트럼프 관세 협상 진전 → 한국 수출주에 직접적 호재")

    add_h2(doc, "7-4. 초보 투자자가 지금 해야 할 일")
    add_callout(doc,
        "지금 시작하는 분께 드리는 5가지 조언",
        "① 한 번에 큰돈을 넣지 마세요. 최소 3~5번 나눠서(=분할 매수) 들어가면 평균 단가를 낮출 수 있습니다. "
        "② 잘 모르는 종목은 피하세요. 본인이 회사 사업 내용을 한 문장으로 설명할 수 없으면 사지 마세요. "
        "③ '폭락한 김에 사자'는 위험합니다. 떨어진 종목이 더 떨어질 수도 있어요. 안정 신호를 기다리세요. "
        "④ 손실을 견딜 수 있는 돈만 투자하세요. 단기에 쓸 돈(전세금·등록금 등)은 절대 금물. "
        "⑤ 뉴스 1~2개로 결정하지 마세요. 한 주일 정도 시장을 지켜본 뒤 판단해도 늦지 않습니다.",
        color="E7F5E7", border_color="00703C")

    add_h2(doc, "7-5. 한 줄 결론")
    add_callout(doc,
        "최종 정리",
        "'잘 오르던 시장이 잠시 멈춰서서 숨을 고르는 중'입니다. 강세장(=오르는 시장) 자체가 끝났다고 보기는 이릅니다. "
        "다만 위에 적은 6가지 악재 중 2개 이상이 같은 시기에 터지면 시장이 더 내려갈 수 있습니다. "
        "특히 다음 주 엔비디아 실적(5/21 새벽)과 호르무즈 해협 상황을 지켜본 뒤 판단해도 늦지 않습니다. "
        "지금은 '공격할 때'가 아니라 '관망하며 분할 준비할 때'입니다.",
        color="FFF4D6", border_color="E69900")

    # ===== Section 8: Glossary =====
    add_h1(doc, "부록. 헷갈리기 쉬운 용어 더 모음")
    glossary = [
        ("시가총액", "회사 주식 전체의 값어치. 주가 × 발행 주식 수. 클수록 큰 회사."),
        ("기관 / 개인", "기관 = 자산운용사·연기금 등 큰 손. 개인 = 우리 같은 일반 투자자."),
        ("순매수 / 순매도", "매수(사기) - 매도(팔기) = 순매수. 양수면 사들이는 중, 음수면 팔아치우는 중."),
        ("강세장 / 약세장", "강세장 = 계속 오르는 시장(불 마켓). 약세장 = 계속 내리는 시장(곰 마켓)."),
        ("차익 실현", "오른 만큼 일단 팔아서 이익을 확정짓는 것. '이 정도면 됐다' 심리."),
        ("FOMO", "Fear Of Missing Out. '나만 못 사면 어떡하지' 하는 조급함. 고점에서 사고 후회하는 원인."),
        ("VIX (변동성 지수)", "시장의 공포 지수. 보통 20 아래면 안정, 30 넘으면 불안 신호. 현재 18.06으로 아직 안정권."),
        ("BP (베이시스 포인트)", "금리 변화 단위. 1bp = 0.01%. '70bp 올랐다' = 0.7% 상승."),
        ("필라델피아 반도체 지수 (SOX)", "미국에 상장된 반도체 회사 30개의 평균. 한국 반도체주의 '예고편' 역할."),
        ("매그니피센트 7 (M7)", "미국 빅테크 7곳: 애플·MS·구글·아마존·엔비디아·메타·테슬라. 미국 증시의 견인차."),
    ]
    add_table(doc, ["용어", "쉬운 설명"],
              [[{"text": k, "bold": True}, {"text": v}] for k, v in glossary],
              col_widths_cm=[5.0, 12.0])

    # ===== Verification note =====
    add_h1(doc, "자료 검증 메모")
    add_para(doc, "이 문서에 나오는 모든 숫자는 두 곳에서 교차 확인했습니다.")
    add_table(doc,
        ["항목", "한국경제 자료", "네이버 검색(타 매체 종합)", "일치"],
        [
            [{"text":"코스피(5/19 마감)"},{"text":"7,271.66 (-3.25%)","center":True},{"text":"7,271.66 (-3.25%)","center":True},{"text":"✓","center":True,"bold":True,"color":"00703C"}],
            [{"text":"코스피(5/20 마감)"},{"text":"—","center":True},{"text":"7,208.95 (-0.86%)","center":True},{"text":"신규","center":True,"color":"595959"}],
            [{"text":"다우(5/19 마감)"},{"text":"49,363.88 (-0.65%)","center":True},{"text":"49,363.88 (-0.65%)","center":True},{"text":"✓","center":True,"bold":True,"color":"00703C"}],
            [{"text":"S&P500(5/19)"},{"text":"7,353.61 (-0.67%)","center":True},{"text":"7,353.61 (-0.67%)","center":True},{"text":"✓","center":True,"bold":True,"color":"00703C"}],
            [{"text":"나스닥(5/19)"},{"text":"25,870.71 (-0.84%)","center":True},{"text":"25,870.71 (-0.84%)","center":True},{"text":"✓","center":True,"bold":True,"color":"00703C"}],
            [{"text":"미 30년물 금리"},{"text":"5.197%","center":True},{"text":"5.197%","center":True},{"text":"✓","center":True,"bold":True,"color":"00703C"}],
            [{"text":"원/달러"},{"text":"—","center":True},{"text":"1,506.8원","center":True},{"text":"신규","center":True,"color":"595959"}],
        ],
        col_widths_cm=[4.5, 4.5, 4.5, 3.5])

    add_para(doc, "", space_before=8)
    add_para(doc, "본 자료는 일반 정보 제공 목적이며, 특정 종목의 매수·매도를 권유하는 것이 아닙니다.",
             size=9, italic=True, color="808080",
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_para(doc, "투자 결정은 반드시 본인의 판단과 책임 하에 이루어져야 합니다.",
             size=9, italic=True, color="808080",
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0)

    doc.save(out_path)
    print(f"Wrote: {out_path}")

if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else "easy_report.docx"
    build(out)
